import os
import gc
from typing import List
from dotenv import load_dotenv

# Suppress ONNX Runtime verbosity warnings in Render logs
os.environ["ORT_LOGGING_LEVEL"] = "3"

load_dotenv()

# --- FastEmbed & Qdrant Vector Store Imports ---
from langchain_community.embeddings import FastEmbedEmbeddings
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams

# --- Text Splitting & Lightweight Document Loaders ---
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader, CSVLoader
import docx

# --- Groq LLM & Vision OCR Imports ---
from groq import Groq

# ----------------- CONFIGURATION & DIRECTORIES -----------------
QDRANT_URL = os.getenv("QDRANT_URL", "")
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY", "")
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
COLLECTION_NAME = "institutional_docs"

# Base folder for local document cache
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOC_FOLDER = os.path.join(BASE_DIR, "documents")
os.makedirs(DOC_FOLDER, exist_ok=True)


# ----------------- EMBEDDINGS & QDRANT CLIENT INITIALIZATION -----------------
def get_embeddings():
    """Initializes FastEmbed BGE-Small embeddings."""
    return FastEmbedEmbeddings(model_name="BAAI/bge-small-en-v1.5")


def get_qdrant_client() -> QdrantClient:
    """Creates a connection client for Qdrant Cloud or local Qdrant instance."""
    if QDRANT_URL and QDRANT_API_KEY:
        return QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, timeout=30.0)
    return QdrantClient(location=":memory:")


def get_vector_store() -> QdrantVectorStore:
    """Ensures Qdrant collection exists and returns the QdrantVectorStore wrapper."""
    client = get_qdrant_client()
    embeddings = get_embeddings()

    # Ensure collection exists
    collections = [col.name for col in client.get_collections().collections]
    if COLLECTION_NAME not in collections:
        client.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(size=384, distance=Distance.COSINE),
        )

    return QdrantVectorStore(
        client=client,
        collection_name=COLLECTION_NAME,
        embedding=embeddings
    )


# ----------------- MULTI-FORMAT DOCUMENT & OCR LOADERS -----------------
def extract_text_from_image(image_path: str) -> str:
    """Uses Groq Vision OCR (llama-3.2-11b-vision-preview) to extract text from images."""
    if not GROQ_API_KEY:
        return "OCR unavailable: Missing GROQ_API_KEY."

    try:
        import base64
        with open(image_path, "rb") as image_file:
            base64_image = base64.b64encode(image_file.read()).decode("utf-8")

        groq_client = Groq(api_key=GROQ_API_KEY)
        completion = groq_client.chat.completions.create(
            model="llama-3.2-11b-vision-preview",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text",
                         "text": "Extract all readable text, tables, notices, and details from this image verbatim."},
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
                        }
                    ]
                }
            ],
            temperature=0.1,
            max_tokens=1024
        )
        return completion.choices[0].message.content or ""
    except Exception as e:
        print(f"⚠️ Groq Vision OCR failed for '{image_path}': {e}")
        return f"Failed to extract text from image: {e}"


def load_document_by_extension(file_path: str) -> List[Document]:
    """Loads documents by file type without requiring heavy C-libraries like unstructured."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif ext in [".docx", ".doc"]:
        doc_obj = docx.Document(file_path)
        text = "\n".join([p.text for p in doc_obj.paragraphs if p.text.strip()])
        return [Document(page_content=text, metadata={"source": os.path.basename(file_path)})]
    elif ext == ".csv":
        loader = CSVLoader(file_path)
        return loader.load()
    elif ext in [".txt", ".md"]:
        loader = TextLoader(file_path, encoding="utf-8")
        return loader.load()
    elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp"]:
        ocr_text = extract_text_from_image(file_path)
        return [Document(page_content=ocr_text, metadata={"source": os.path.basename(file_path)})]
    else:
        raise ValueError(f"Unsupported file extension: {ext}")


# ----------------- MEMORY-SAFE BATCHED INDEXING -----------------
def index_single_file(file_path: str, filename: str) -> int:
    """
    Loads, chunks, and indexes files into Qdrant Cloud in 5-chunk memory batches
    to prevent RAM spikes/OOM crashes on Render Free Tier.
    """
    try:
        print(f"⏳ Loading and parsing '{filename}'...")
        raw_docs = load_document_by_extension(file_path)

        # Chunk text
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        chunks = text_splitter.split_documents(raw_docs)

        if not chunks:
            print(f"⚠️ No text extracted from '{filename}'.")
            return 0

        # Inject source metadata into every chunk
        for i, chunk in enumerate(chunks):
            chunk.metadata["source"] = filename
            chunk.metadata["chunk_id"] = i

        # Upsert vectors in small batches of 5 (Prevents 502 OOM restarts)
        vector_store = get_vector_store()
        batch_size = 5
        total_chunks = len(chunks)

        for i in range(0, total_chunks, batch_size):
            batch = chunks[i:i + batch_size]
            vector_store.add_documents(batch)
            print(f"   Indexed batch {min(i + batch_size, total_chunks)}/{total_chunks} chunks for '{filename}'...")

        # Explicit memory cleanup
        del raw_docs, chunks
        gc.collect()

        print(f"✅ Successfully indexed '{filename}' ({total_chunks} chunks total).")
        return total_chunks

    except Exception as e:
        print(f"❌ Error indexing '{filename}': {e}")
        gc.collect()
        raise e


# ----------------- RAG QUERY PIPELINE -----------------
def query_rag_system(user_query: str) -> str:
    """Performs vector similarity search on Qdrant and generates an answer using Groq LLM."""
    if not user_query.strip():
        return "Please ask a valid question."

    try:
        # 1. Similarity search on Qdrant Cloud
        vector_store = get_vector_store()
        docs = vector_store.similarity_search(user_query, k=4)

        if not docs:
            return "This information is not available in the college records."

        # 2. Build context block
        context_parts = []
        for d in docs:
            src = d.metadata.get("source", "Institutional Document")
            context_parts.append(f"--- Document Source: {src} ---\n{d.page_content}")

        context_text = "\n\n".join(context_parts)

        # 3. Generate response using Groq
        if not GROQ_API_KEY:
            return f"Retrieved Context:\n{context_text}\n\n(LLM generation offline: Missing GROQ_API_KEY)"

        groq_client = Groq(api_key=GROQ_API_KEY)

        system_prompt = (
            "You are the official AI Academic Assistant for SNJB College of Engineering.\n"
            "Answer student queries accurately using strictly the provided context.\n"
            "If the context does not contain the answer, state clearly: "
            "'This information is not available in the college records.'\n"
            "Keep responses concise, polite, and well-formatted."
        )

        user_prompt = f"Context Information:\n{context_text}\n\nStudent Question: {user_query}"

        completion = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.2,
            max_tokens=800
        )

        answer = completion.choices[0].message.content or "No answer generated."

        # Clean up local memory
        del docs
        gc.collect()

        return answer

    except Exception as e:
        print(f"❌ Error in query_rag_system: {e}")
        gc.collect()
        return f"An error occurred while retrieving answers: {str(e)}"